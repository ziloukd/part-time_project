
import os
import json
from docx import Document
import requests
from lxml import etree
import re
import time


CITY_LIST = ['武汉','宜昌','黄石','十堰','襄阳','鄂州','荆州','荆门','黄冈','咸宁','孝感','随州','恩施','神农架','潜江', '天门', '仙桃']
headers = {'User-Agent':'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/81.0.4044.138 Safari/537.36'}



def get_detail(url,data_dic,path):
    for retry_time in range(3):
        try:
            response = requests.get(url, headers=headers, timeout = 10)
            if response.status_code == 200:
                break
            time.sleep(2)
        except:
            pass
    # 解析网页完善数据
    # 锁定抓取内容的大致位置
    ret = re.compile('<!-- 行程安排.*?-->(.*?)<!-- 行程安排.*?-->', re.S)
    html = ret.findall(response.text)[0]
    tree = etree.HTML(html)
    data_dic['游玩时间'] = data_dic['发表时间']
    try:
        date_tree = etree.HTML(response.text)
        data_time = date_tree.xpath("//li[@class='date']/span[1]/text()")[0]
        data_dic['发表时间'] = re.sub(r'/','-',data_time)

    except:
        pass
    try:
        data_dic['天数'] = tree.xpath("//li//p[contains(text(), '天数')]/span[2]/text()")[0]
    except:
        pass
    try:
        play_list = filter(lambda d:d if not re.findall('\s',d) else '', tree.xpath("//li//p[contains(text(), '玩法')]//text()"))
        play = ','.join(i.strip() for i in play_list)
        data_dic['玩法'] = re.sub('玩法,/,', '', play)
    except:
        pass
    try:
        cost = ' '.join(i.strip() for i in tree.xpath("//li//p[contains(text(), '人均费用')]//text()"))
        data_dic['人均花费'] = re.sub('人均费用 / ','',cost)

    except:
        pass
    try:
        people_list = filter(lambda d:d if not re.findall('\s',d) else '', tree.xpath("//li//p[contains(text(), '人物')]//text()"))
        people = ','.join(i.strip() for i in people_list)
        data_dic['和谁'] = re.sub('人物,/,','',people)
    except:
        pass

    try:
        content_list = tree.xpath("//div[@id='b_panel_schedule']//text()")
        content = re.sub(r'\+1', '', ''.join([i.strip() for i in content_list]))
        data_dic['正文'] = content
    except:
        pass
    try:
        src_list = tree.xpath("//img/@data-original")
        data_dic['src_url'] = src_list
    except:
        pass
    ret = re.compile('<!-- 最新评论.*?-->(.*?)<!-- 最新评论.*?-->', re.S)
    html = ret.findall(response.text)[0]
    comment_tree = etree.HTML(html)
    try:
        comment_list = comment_tree.xpath("//dl/dd/div/text()")
        comment = '\n'.join([str(index + 1) + '.' + i.strip() for index, i in enumerate(comment_list)])
        data_dic['评论'] = comment
    except:
        pass

    with open(path, 'r', encoding='utf-8') as f1:
        data_list = json.load(f1)
    data_list.append(data_dic)
    with open(path, 'w', encoding='utf-8') as f1:
        f1.write(json.dumps(data_list))
    print(data_dic)

def sort_func(dic):
    # 日期转时间戳
    timeArray = time.strptime(dic['发表时间'], "%Y-%m-%d")
    # 轉換為時間戳:
    timeStamp = int(time.mktime(timeArray))
    return timeStamp


for year in ('2017','2018'):
    for orgin in ('途牛','去哪儿','携程'):
        if orgin in ('途牛','携程'):
            continue
        for city in CITY_LIST:
            doc = Document()
            file_catalog = '../%s/%s/%s' %(year,orgin,city)
            file_name = '%s.json' % city
            file_name_new = '%s_new.json' % city
            file_name_doc = '%s.docx' % city
            path = os.path.join(file_catalog, file_name)
            path_new = os.path.join(file_catalog, file_name_new)
            save_path = os.path.join(file_catalog, file_name_doc)

            # 文档保存路径
            bak_catalog = '../%s/%s/%s' %(year,orgin,city)
            if not os.path.exists(bak_catalog):
                os.makedirs(bak_catalog)

            bak_save_path = os.path.join(bak_catalog,file_name_doc)
            with open(path_new, encoding='utf-8') as f1:
                li = json.load(f1)
                # 排序
            temp_list = list(set([str(i) for i in li]))
            li = [eval(i) for i in temp_list]
            if li:

                data_list = sorted(li, key=sort_func)
                for data_num,data in enumerate(data_list):
                    # 写入
                    # get_detail(data['网址'], data, path)
                    # time.sleep(3)
                    doc.add_heading(f"{data_num+1}.{data['标题']}", 2)
                    doc.add_paragraph(data['网址'])
                    doc.add_paragraph('来源：%s' % orgin)
                    doc.add_paragraph(f"发表时间：{data['发表时间']}")
                    doc.add_paragraph(f"天数：{data['天数']}")
                    doc.add_paragraph(f"游玩时间：{data['游玩时间']}")
                    try:
                        doc.add_paragraph(f"人均花费：{data['人均花费']}")
                    except:
                        doc.add_paragraph(f"人均花费：{data['人均']}")
                    doc.add_paragraph(f"和谁：{data['和谁']}")
                    doc.add_paragraph(f"玩法：{data['玩法']}")
                    doc.add_paragraph(f"旅游路线：{data['旅游路线']}")
                    doc.add_paragraph(f"正文：\n{data['正文']}")
                    doc.add_paragraph(f"评论：\n{data['评论']}")
            print('=' * 50)
            print(bak_save_path)
            doc.save(bak_save_path)